import os
import re
import shutil
import platform
from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document
import pytesseract
from PIL import Image
from typing import Optional
from pathlib import Path
import logging

logger = logging.getLogger(__name__)


class FileProcessor:
    def __init__(self):
        # Configure Tesseract path based on OS
        self.tesseract_path = self._find_tesseract()
        if self.tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = self.tesseract_path
            logger.info(f"Tesseract configured at: {self.tesseract_path}")
        else:
            logger.warning("Tesseract not found - OCR for images will be unavailable")
    
    def _find_tesseract(self) -> Optional[str]:
        """Find Tesseract executable based on operating system."""
        system = platform.system()
        
        # Common paths by OS
        paths_to_check = []
        
        if system == "Windows":
            paths_to_check = [
                r'C:\Program Files\Tesseract-OCR\tesseract.exe',
                r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
                r'C:\Tesseract-OCR\tesseract.exe',
            ]
        elif system == "Darwin":  # macOS
            paths_to_check = [
                '/usr/local/bin/tesseract',
                '/opt/homebrew/bin/tesseract',  # Apple Silicon
                '/opt/local/bin/tesseract',  # MacPorts
            ]
        else:  # Linux
            paths_to_check = [
                '/usr/bin/tesseract',
                '/usr/local/bin/tesseract',
            ]
        
        # Check explicit paths first
        for path in paths_to_check:
            if os.path.isfile(path):
                return path
        
        # Try to find in PATH
        tesseract_in_path = shutil.which('tesseract')
        if tesseract_in_path:
            return tesseract_in_path
        
        return None
    
    def is_ocr_available(self) -> bool:
        """Check if OCR functionality is available."""
        return self.tesseract_path is not None

    def _detect_file_type(self, file_path: str) -> str:
        """Detect file type from magic bytes if extension is missing"""
        try:
            with open(file_path, 'rb') as f:
                header = f.read(8)
            
            # PDF: starts with %PDF
            if header[:4] == b'%PDF':
                return '.pdf'
            
            # DOCX/ZIP: starts with PK (ZIP archive)
            if header[:2] == b'PK':
                return '.docx'
            
            # PNG: starts with \x89PNG
            if header[:4] == b'\x89PNG':
                return '.png'
            
            # JPEG: starts with \xff\xd8\xff
            if header[:3] == b'\xff\xd8\xff':
                return '.jpg'
            
            # Default to txt if we can't detect
            return '.txt'
        except:
            return '.txt'

    def extract_text(self, file_path: str) -> str:
        """Extract text from various file formats"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_ext = Path(file_path).suffix.lower()
        
        # Always detect file type from content (magic bytes) for more reliable detection
        # File extensions can be wrong or missing
        detected_ext = self._detect_file_type(file_path)
        
        logger.info(f"File extraction: path={file_path}, extension={file_ext}, detected={detected_ext}")
        
        # Use detected type if extension is empty/missing, or if detected type differs
        # (trust magic bytes over file extension for reliability)
        if not file_ext or file_ext == '.' or detected_ext != '.txt':
            # If we detected a real file type, use it
            if detected_ext in ['.pdf', '.docx', '.png', '.jpg']:
                file_ext = detected_ext
                logger.info(f"Using detected file type: {file_ext}")
        
        if file_ext == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_ext == '.docx':
            return self._extract_from_docx(file_path)
        elif file_ext in ('.png', '.jpg', '.jpeg'):
            return self._extract_from_image(file_path)
        elif file_ext == '.txt':
            return self._extract_from_txt(file_path)
        elif file_ext == '.doc':
            # Try to handle .doc as .docx (may not always work)
            try:
                return self._extract_from_docx(file_path)
            except:
                raise ValueError(f"Old .doc format not supported. Please convert to .docx or .pdf")
        else:
            # Try PDF as default (most common resume format)
            try:
                return self._extract_from_pdf(file_path)
            except:
                raise ValueError(f"Unsupported file format: {file_ext}")

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        try:
            text = pdf_extract_text(file_path)
            return self._clean_text(text)
        except Exception as e:
            raise ValueError(f"PDF extraction failed: {str(e)}")

    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files - includes paragraphs and tables"""
        try:
            doc = Document(file_path)
            full_text = []
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Also extract from tables (resumes often use tables for layout)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text and cell_text not in row_text:  # Avoid duplicates
                            row_text.append(cell_text)
                    if row_text:
                        full_text.append(' | '.join(row_text))
            
            return '\n'.join(full_text)
        except Exception as e:
            raise ValueError(f"DOCX extraction failed: {str(e)}")

    def _extract_from_image(self, file_path: str) -> str:
        """Extract text from image files using OCR"""
        if not self.is_ocr_available():
            raise ValueError(
                "OCR is not available. Please install Tesseract: "
                "macOS: 'brew install tesseract', "
                "Ubuntu: 'sudo apt-get install tesseract-ocr', "
                "Windows: Download from https://github.com/UB-Mannheim/tesseract/wiki"
            )
        try:
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img)
            return self._clean_text(text)
        except Exception as e:
            raise ValueError(f"Image OCR failed: {str(e)}")

    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            raise ValueError(f"Text file reading failed: {str(e)}")

    def _clean_text(self, text: str) -> str:
        """Clean extracted text by removing excessive whitespace and special characters"""
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        # Replace multiple spaces (not newlines) with single space
        text = re.sub(r'[^\S\n]+', ' ', text)
        # Replace 3+ consecutive newlines with 2 newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        # Remove form feed characters
        text = re.sub(r'\x0c', '', text)
        # Remove non-printable characters (except newlines, tabs)
        text = ''.join(char for char in text if char.isprintable() or char in {'\n', '\t'})
        return text.strip()